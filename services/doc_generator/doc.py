import re
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

from data_structures.story import (
    StoryData,
)  # Adjust the import based on your project structure

logger = logging.getLogger(__name__)


class StoryDocument:
    def __init__(self, story_data: StoryData) -> None:
        self.chapters = story_data.post_processed_chapters
        self.image_filenames = story_data.image_filenames
        self.story_folder = story_data.story_folder
        self.images_folder = story_data.images_subfolder
        self.title = story_data.title
        self.doc = Document()
        logger.info("StoryDocument initialized.")

    def _add_stylized_title(self, text: str, level: int):
        """Adds a stylized title or heading to the document."""
        heading = self.doc.add_heading("", level=level)
        # Create a new run if no runs exist
        if len(heading.runs) == 0:
            run = heading.add_run()
        else:
            run = heading.runs[0]
        run.text = text
        font = run.font
        font.name = "Times New Roman"
        font.size = Pt(24 if level == 0 else 18)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")

    def _process_paragraph(self, paragraph: str, chapter_number: int) -> None:
        """Processes a single paragraph to add it to the final document and embed images."""
        # Keep track of the starting point in the paragraph
        start = 0
        for match in re.finditer(r"\[img: (\d+)\.(\d+)\]", paragraph):
            # Add text from the last match up to this one
            pre_image_text = paragraph[start : match.start()].strip()
            if pre_image_text:
                self.doc.add_paragraph(pre_image_text)

            # Add the image
            chapter_num, image_index = map(int, match.groups())
            if chapter_num == chapter_number:
                self._add_image(chapter_number, image_index)

            # Update the starting point for the next loop
            start = match.end()

        # Add any remaining text after the last image
        post_image_text = paragraph[start:].strip()
        if post_image_text:
            self.doc.add_paragraph(post_image_text)

    def _add_image(self, chapter_number: int, image_prompt_index: int) -> None:
        """Adds an image to the document based on the chapter number and image index."""
        filename = f"Chapter_{chapter_number}_Image_{image_prompt_index}.png"
        filepath = os.path.join(self.story_folder, self.images_folder, filename)
        if os.path.exists(filepath):
            para = self.doc.add_paragraph()
            run = para.add_run()
            run.add_picture(filepath, width=Inches(4))
        else:
            logger.warning(f"Image {filename} not found. Skipping image insertion.")

    def create_document(self) -> None:
        """Create the story document by adding titles and processing each chapter and paragraph."""
        self._add_stylized_title(self.title, 0)
        for chapter_number, chapter in enumerate(self.chapters, start=1):
            self._add_stylized_title(chapter["title"], 1)
            for paragraph in chapter["content"].split("\n"):
                self._process_paragraph(paragraph, chapter_number)

    def save_document(self, filename: str) -> str:
        """Save the generated story document to a specified file."""
        filepath = os.path.join(self.story_folder, filename)
        self.doc.save(filepath)
        logging.info(f"Story document saved in: {filepath}")
        return filepath
